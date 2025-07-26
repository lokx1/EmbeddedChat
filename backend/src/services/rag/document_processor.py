import os
import io
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document as DocxDocument
import openpyxl
from ...utils.logger import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Process various document types and extract text"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        
            return {
                "success": True,
                "text": text.strip(),
                "page_count": len(pdf_reader.pages)
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " | ".join(row_text) + "\n"
                    
            return {
                "success": True,
                "text": text.strip(),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def extract_text_from_xlsx(file_path: str) -> Dict[str, Any]:
        """Extract text from XLSX file"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""
            sheet_count = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n--- Sheet: {sheet_name} ---\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                
                sheet_count += 1
                
            return {
                "success": True,
                "text": text.strip(),
                "sheet_count": sheet_count
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from XLSX {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                
            return {
                "success": True,
                "text": text.strip(),
                "line_count": len(text.split('\n'))
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    return {
                        "success": True,
                        "text": text.strip(),
                        "line_count": len(text.split('\n')),
                        "encoding": encoding
                    }
                except:
                    continue
                    
            logger.error(f"Could not decode text file {file_path}")
            return {"success": False, "error": "Could not decode text file"}
            
        except Exception as e:
            logger.error(f"Failed to extract text from TXT {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    @classmethod
    def extract_text(cls, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Extract text based on file type"""
        try:
            if mime_type == 'application/pdf':
                return cls.extract_text_from_pdf(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return cls.extract_text_from_docx(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
                return cls.extract_text_from_xlsx(file_path)
            elif mime_type.startswith('text/'):
                return cls.extract_text_from_txt(file_path)
            else:
                return {"success": False, "error": f"Unsupported file type: {mime_type}"}
                
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    @staticmethod
    def generate_summary(text: str, max_length: int = 500) -> str:
        """Generate a simple summary of the text"""
        if len(text) <= max_length:
            return text
            
        # Simple extractive summary - take first few sentences
        sentences = text.split('. ')
        summary = ""
        for sentence in sentences:
            if len(summary + sentence) <= max_length:
                summary += sentence + ". "
            else:
                break
                
        return summary.strip() or text[:max_length] + "..."


# Global document processor instance
document_processor = DocumentProcessor()
